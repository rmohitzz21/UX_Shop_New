from pathlib import Path
from datetime import date
from textwrap import wrap

from PIL import Image, ImageDraw, ImageFont
from docx import Document
from docx.enum.section import WD_SECTION
from docx.enum.table import WD_ALIGN_VERTICAL, WD_TABLE_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_BREAK
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Inches, Pt, RGBColor


ROOT = Path(__file__).resolve().parents[1]
OUT = ROOT / "deliverables"
ASSETS = OUT / "report_assets"
OUT.mkdir(exist_ok=True)
ASSETS.mkdir(exist_ok=True)
DOCX_PATH = OUT / "UX_Pacific_Shop_Master_Project_Report_2026-06-01.docx"

PURPLE = "6F4BFF"
PURPLE_LIGHT = "EEEAFE"
DARK = "111827"
MUTED = "5B6475"
GREEN = "0F766E"
RED = "B42318"
AMBER = "B54708"
BLUE = "175CD3"
LIGHT_GREY = "F2F4F7"
BORDER = "D0D5DD"


def font(size=24, bold=False):
    candidates = [
        "C:/Windows/Fonts/arialbd.ttf" if bold else "C:/Windows/Fonts/arial.ttf",
        "C:/Windows/Fonts/calibrib.ttf" if bold else "C:/Windows/Fonts/calibri.ttf",
    ]
    for candidate in candidates:
        if Path(candidate).exists():
            return ImageFont.truetype(candidate, size)
    return ImageFont.load_default()


def wrapped_lines(text, width):
    return wrap(text, width=width, break_long_words=False, replace_whitespace=False) or [""]


def box(draw, xy, title, body="", fill="#F8FAFC", outline="#C7D2FE", title_fill="#4F46E5"):
    x1, y1, x2, y2 = xy
    draw.rounded_rectangle(xy, radius=16, fill=fill, outline=outline, width=3)
    draw.rounded_rectangle((x1, y1, x2, y1 + 42), radius=16, fill=title_fill, outline=title_fill)
    draw.rectangle((x1, y1 + 24, x2, y1 + 42), fill=title_fill)
    draw.text((x1 + 15, y1 + 10), title, font=font(20, True), fill="white")
    y = y1 + 56
    for line in body.split("\n"):
        for fragment in wrapped_lines(line, 31):
            draw.text((x1 + 16, y), fragment, font=font(17), fill="#1F2937")
            y += 23


def arrow(draw, start, end, label=""):
    draw.line((start, end), fill="#6F4BFF", width=4)
    x2, y2 = end
    x1, y1 = start
    if abs(x2 - x1) >= abs(y2 - y1):
        sign = 1 if x2 > x1 else -1
        points = [(x2, y2), (x2 - sign * 16, y2 - 9), (x2 - sign * 16, y2 + 9)]
    else:
        sign = 1 if y2 > y1 else -1
        points = [(x2, y2), (x2 - 9, y2 - sign * 16), (x2 + 9, y2 - sign * 16)]
    draw.polygon(points, fill="#6F4BFF")
    if label:
        lx = (x1 + x2) // 2
        ly = (y1 + y2) // 2 - 22
        draw.rounded_rectangle((lx - 65, ly - 3, lx + 65, ly + 24), radius=8, fill="#FFFFFF")
        draw.text((lx - 58, ly), label, font=font(15, True), fill="#4F46E5")


def save_diagram(name, title, subtitle, boxes, arrows):
    image = Image.new("RGB", (1600, 980), "white")
    draw = ImageDraw.Draw(image)
    draw.text((60, 42), title, font=font(38, True), fill="#111827")
    draw.text((60, 94), subtitle, font=font(20), fill="#475467")
    for arrow_spec in arrows:
        arrow(draw, *arrow_spec)
    for box_spec in boxes:
        box(draw, *box_spec)
    draw.text((60, 934), "UX Pacific Shop | Master Project Report", font=font(16), fill="#667085")
    path = ASSETS / name
    image.save(path, quality=95)
    return path


def build_diagrams():
    ia = save_diagram(
        "information_architecture.png",
        "Information Architecture",
        "Customer storefront, authenticated area, and administration workspace",
        [
            ((620, 155, 980, 265), "UX Pacific Shop", "Digital + physical UX resources"),
            ((90, 350, 400, 610), "Discover", "Home\nShop All\nCategories\nBundles\nFreebies\nSearch\nProduct detail"),
            ((460, 350, 770, 610), "Purchase", "Cart\nSign in / Sign up\nCheckout\nPayment\nOrder confirmation"),
            ((830, 350, 1140, 610), "Customer", "Account\nProfile\nAddresses\nOrders\nOrder detail\nWishlist\nReviews"),
            ((1200, 350, 1510, 610), "Admin", "Dashboard\nCatalog\nCategories\nBundles\nFreebies\nOrders\nUsers\nReviews\nMessages"),
            ((460, 720, 770, 865), "Support", "Contact\nPolicies\n404 / 500\nEmail notifications"),
            ((830, 720, 1140, 865), "Growth", "Featured items\nCoupons - required\nNewsletter - required\nAnalytics - required"),
        ],
        [
            ((800, 265), (245, 350), "browse"),
            ((800, 265), (615, 350), "buy"),
            ((800, 265), (985, 350), "manage"),
            ((800, 265), (1355, 350), "operate"),
            ((615, 610), (615, 720), "help"),
            ((985, 610), (985, 720), "improve"),
        ],
    )
    purchase = save_diagram(
        "purchase_flow.png",
        "Customer Purchase Flow",
        "Shared checkout with a controlled branch for digital and physical fulfilment",
        [
            ((70, 200, 310, 330), "1. Discover", "Browse, search,\nfilter, open PDP"),
            ((390, 200, 630, 330), "2. Configure", "Choose license/type,\nquantity, add to cart"),
            ((710, 200, 950, 330), "3. Checkout", "Authenticate,\nvalidate cart"),
            ((1030, 200, 1270, 330), "4. Pay", "COD or Razorpay\nCard / UPI"),
            ((1350, 200, 1570, 330), "5. Confirm", "Order created,\nemail sent"),
            ((900, 540, 1170, 700), "Digital Order", "Email delivery\nDownload entitlement\nNo shipping charge\nCOD unavailable"),
            ((1250, 540, 1520, 700), "Physical Order", "Shipping address\nINR 50 shipping rule\nInventory decrement\nTracking lifecycle"),
            ((250, 540, 600, 700), "Customer Follow-up", "Order history\nOrder detail\nReview eligibility\nSupport contact"),
        ],
        [
            ((310, 265), (390, 265), ""),
            ((630, 265), (710, 265), ""),
            ((950, 265), (1030, 265), ""),
            ((1270, 265), (1350, 265), ""),
            ((1460, 330), (1035, 540), "digital"),
            ((1460, 330), (1385, 540), "physical"),
            ((1035, 700), (600, 620), "account"),
            ((1385, 700), (600, 650), "account"),
        ],
    )
    architecture = save_diagram(
        "system_architecture.png",
        "Current System Architecture",
        "Server-rendered PHP storefront with JSON APIs, MySQL persistence, email, and Razorpay",
        [
            ((90, 170, 360, 330), "Browser", "Desktop / mobile\nPHP pages\nJavaScript fetch calls"),
            ((480, 170, 770, 330), "Apache + PHP", "Public pages\nAuth guards\nAdmin pages\n.htaccess rules"),
            ((900, 170, 1210, 330), "API Layer", "auth/*\ncart/*\ncatalog/*\norder/*\npayment/*\nadmin/*"),
            ((900, 540, 1210, 700), "MySQL", "Catalog\nUsers\nCart\nOrders\nAdmin\nReviews"),
            ((1280, 460, 1530, 610), "Razorpay", "Create order\nVerify payment\nWebhook"),
            ((1280, 690, 1530, 840), "SMTP", "Transactional email\nPassword reset\nOrder email"),
            ((480, 540, 770, 700), "Static Assets", "CSS\nJavaScript\nImages\nProduct files"),
        ],
        [
            ((360, 250), (480, 250), "request"),
            ((770, 250), (900, 250), "fetch"),
            ((1055, 330), (1055, 540), "SQL"),
            ((1210, 520), (1280, 520), "payment"),
            ((1210, 640), (1280, 760), "email"),
            ((625, 330), (625, 540), "serve"),
        ],
    )
    lifecycle = save_diagram(
        "order_lifecycle.png",
        "Order Lifecycle",
        "Recommended state model for payment, fulfilment, customer visibility, and admin control",
        [
            ((80, 250, 330, 390), "Cart", "Validated product IDs\nServer-side price source"),
            ((410, 250, 660, 390), "Draft Order", "Awaiting payment\nReserve stock only"),
            ((740, 150, 990, 290), "Paid", "Payment verified\nCommit stock movement"),
            ((740, 390, 990, 530), "COD Confirmed", "Physical only\nCommit stock movement"),
            ((1070, 250, 1320, 390), "Processing", "Prepare download\nor parcel"),
            ((1240, 540, 1510, 680), "Delivered", "Download available\nor parcel delivered"),
            ((410, 610, 660, 750), "Failed / Cancelled", "Release reservation\nPreserve cart\nRecord reason"),
        ],
        [
            ((330, 320), (410, 320), "checkout"),
            ((660, 310), (740, 220), "online"),
            ((660, 350), (740, 460), "COD"),
            ((990, 220), (1070, 300), ""),
            ((990, 460), (1070, 340), ""),
            ((1320, 390), (1375, 540), "fulfil"),
            ((535, 390), (535, 610), "fail"),
        ],
    )
    return ia, purchase, architecture, lifecycle


def set_cell_shading(cell, fill):
    tc_pr = cell._tc.get_or_add_tcPr()
    shd = tc_pr.find(qn("w:shd"))
    if shd is None:
        shd = OxmlElement("w:shd")
        tc_pr.append(shd)
    shd.set(qn("w:fill"), fill)


def set_cell_text(cell, text, color=DARK, bold=False, size=8.6):
    cell.text = ""
    paragraph = cell.paragraphs[0]
    paragraph.paragraph_format.space_after = Pt(0)
    run = paragraph.add_run(str(text))
    run.font.name = "Calibri"
    run.font.size = Pt(size)
    run.font.bold = bold
    run.font.color.rgb = RGBColor.from_string(color)
    cell.vertical_alignment = WD_ALIGN_VERTICAL.TOP


def set_cell_margins(cell, top=70, start=90, bottom=70, end=90):
    tc = cell._tc
    tc_pr = tc.get_or_add_tcPr()
    tc_mar = tc_pr.first_child_found_in("w:tcMar")
    if tc_mar is None:
        tc_mar = OxmlElement("w:tcMar")
        tc_pr.append(tc_mar)
    for margin, value in [("top", top), ("start", start), ("bottom", bottom), ("end", end)]:
        node = tc_mar.find(qn(f"w:{margin}"))
        if node is None:
            node = OxmlElement(f"w:{margin}")
            tc_mar.append(node)
        node.set(qn("w:w"), str(value))
        node.set(qn("w:type"), "dxa")


def repeat_header(row):
    tr_pr = row._tr.get_or_add_trPr()
    tbl_header = OxmlElement("w:tblHeader")
    tbl_header.set(qn("w:val"), "true")
    tr_pr.append(tbl_header)


def cant_split(row):
    tr_pr = row._tr.get_or_add_trPr()
    node = OxmlElement("w:cantSplit")
    tr_pr.append(node)


def set_repeat_table_header(table):
    repeat_header(table.rows[0])
    for row in table.rows:
        cant_split(row)


def set_col_widths(table, widths):
    for row in table.rows:
        for idx, width in enumerate(widths):
            row.cells[idx].width = Inches(width)


def add_table(doc, headers, rows, widths=None, font_size=8.4):
    table = doc.add_table(rows=1, cols=len(headers))
    table.style = "Table Grid"
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    table.autofit = False
    for idx, header in enumerate(headers):
        set_cell_shading(table.rows[0].cells[idx], PURPLE)
        set_cell_text(table.rows[0].cells[idx], header, color="FFFFFF", bold=True, size=8.7)
    for row_index, row_data in enumerate(rows):
        cells = table.add_row().cells
        fill = "FFFFFF" if row_index % 2 == 0 else "F8FAFC"
        for idx, value in enumerate(row_data):
            set_cell_shading(cells[idx], fill)
            set_cell_margins(cells[idx])
            set_cell_text(cells[idx], value, size=font_size)
    if widths:
        set_col_widths(table, widths)
    set_repeat_table_header(table)
    doc.add_paragraph("")
    return table


def add_page_field(paragraph, instruction):
    run = paragraph.add_run()
    fld_char_1 = OxmlElement("w:fldChar")
    fld_char_1.set(qn("w:fldCharType"), "begin")
    instr_text = OxmlElement("w:instrText")
    instr_text.set(qn("xml:space"), "preserve")
    instr_text.text = instruction
    fld_char_2 = OxmlElement("w:fldChar")
    fld_char_2.set(qn("w:fldCharType"), "end")
    run._r.append(fld_char_1)
    run._r.append(instr_text)
    run._r.append(fld_char_2)


def add_heading(doc, text, level=1):
    paragraph = doc.add_heading(text, level)
    paragraph.paragraph_format.keep_with_next = True
    return paragraph


def add_bullets(doc, items, style="List Bullet"):
    for item in items:
        paragraph = doc.add_paragraph(style=style)
        paragraph.paragraph_format.space_after = Pt(2)
        paragraph.add_run(item)


def add_callout(doc, title, body, color=PURPLE_LIGHT):
    table = doc.add_table(rows=1, cols=1)
    table.autofit = False
    cell = table.cell(0, 0)
    set_cell_shading(cell, color)
    set_cell_margins(cell, top=120, start=160, bottom=120, end=160)
    paragraph = cell.paragraphs[0]
    paragraph.paragraph_format.space_after = Pt(3)
    run = paragraph.add_run(title)
    run.bold = True
    run.font.color.rgb = RGBColor.from_string(PURPLE)
    paragraph = cell.add_paragraph(body)
    paragraph.paragraph_format.space_after = Pt(0)
    paragraph.style = "Body Text"
    doc.add_paragraph("")


def add_status_key(doc):
    add_table(
        doc,
        ["Status", "Meaning"],
        [
            ("Implemented", "Present in the current codebase or database and available for verification."),
            ("Partial", "Present, but incomplete, inconsistent, or blocked by a defect."),
            ("Required", "Recommended minimum for a credible production launch."),
            ("Future", "Useful growth capability after the launch baseline is stable."),
        ],
        [1.25, 5.45],
    )


def setup_document():
    doc = Document()
    section = doc.sections[0]
    section.top_margin = Inches(0.8)
    section.bottom_margin = Inches(0.75)
    section.left_margin = Inches(0.85)
    section.right_margin = Inches(0.85)

    styles = doc.styles
    normal = styles["Normal"]
    normal.font.name = "Calibri"
    normal.font.size = Pt(10)
    normal.font.color.rgb = RGBColor.from_string(DARK)
    normal.paragraph_format.space_after = Pt(4)
    normal.paragraph_format.line_spacing = 1.05

    styles["Body Text"].font.name = "Calibri"
    styles["Body Text"].font.size = Pt(10)
    styles["Body Text"].font.color.rgb = RGBColor.from_string(DARK)
    styles["Body Text"].paragraph_format.space_after = Pt(4)

    for name, size, color in [
        ("Title", 30, DARK),
        ("Subtitle", 15, MUTED),
        ("Heading 1", 17, PURPLE),
        ("Heading 2", 13, DARK),
        ("Heading 3", 11, PURPLE),
    ]:
        style = styles[name]
        style.font.name = "Calibri"
        style.font.size = Pt(size)
        style.font.color.rgb = RGBColor.from_string(color)
        style.font.bold = True
        style.paragraph_format.space_before = Pt(9)
        style.paragraph_format.space_after = Pt(4)
        style.paragraph_format.keep_with_next = True

    for name in ["List Bullet", "List Number"]:
        styles[name].font.name = "Calibri"
        styles[name].font.size = Pt(9.7)
        styles[name].paragraph_format.space_after = Pt(2)

    header = section.header
    paragraph = header.paragraphs[0]
    paragraph.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    run = paragraph.add_run("UX PACIFIC SHOP  |  MASTER PROJECT REPORT")
    run.font.name = "Calibri"
    run.font.size = Pt(8)
    run.font.color.rgb = RGBColor.from_string(MUTED)

    footer = section.footer
    paragraph = footer.paragraphs[0]
    paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = paragraph.add_run("CONFIDENTIAL PROJECT BLUEPRINT  |  ")
    run.font.name = "Calibri"
    run.font.size = Pt(8)
    run.font.color.rgb = RGBColor.from_string(MUTED)
    add_page_field(paragraph, "PAGE")
    paragraph.add_run(" / ")
    add_page_field(paragraph, "NUMPAGES")
    return doc


def cover(doc):
    paragraph = doc.add_paragraph()
    paragraph.paragraph_format.space_before = Pt(45)
    run = paragraph.add_run("UX PACIFIC")
    run.font.name = "Calibri"
    run.font.size = Pt(16)
    run.font.bold = True
    run.font.color.rgb = RGBColor.from_string(PURPLE)

    paragraph = doc.add_paragraph()
    paragraph.paragraph_format.space_before = Pt(60)
    run = paragraph.add_run("UX Pacific Shop")
    run.font.name = "Calibri"
    run.font.size = Pt(36)
    run.font.bold = True
    run.font.color.rgb = RGBColor.from_string(DARK)

    paragraph = doc.add_paragraph()
    run = paragraph.add_run("Master Ecommerce Product, Feature, Architecture,\nDesign System and Launch Blueprint")
    run.font.name = "Calibri"
    run.font.size = Pt(19)
    run.font.bold = True
    run.font.color.rgb = RGBColor.from_string(PURPLE)

    paragraph = doc.add_paragraph()
    paragraph.paragraph_format.space_before = Pt(18)
    run = paragraph.add_run(
        "A professional reference for strategy, product scope, digital and physical commerce, "
        "customer journeys, administration, URLs, APIs, pricing, UI rules, QA risks and implementation priorities."
    )
    run.font.name = "Calibri"
    run.font.size = Pt(12)
    run.font.color.rgb = RGBColor.from_string(MUTED)

    doc.add_paragraph("")
    table = doc.add_table(rows=4, cols=2)
    table.autofit = False
    table.style = "Table Grid"
    fields = [
        ("Prepared for", "UX Pacific Shop"),
        ("Prepared on", "01 June 2026"),
        ("Document type", "Master project report and production implementation blueprint"),
        ("Current source", r"C:\xampp\htdocs\SHOPMohitSir\UX_Shop_New"),
    ]
    for row, values in zip(table.rows, fields):
        set_cell_shading(row.cells[0], PURPLE_LIGHT)
        set_cell_text(row.cells[0], values[0], color=PURPLE, bold=True, size=9)
        set_cell_text(row.cells[1], values[1], size=9)
        for cell in row.cells:
            set_cell_margins(cell, top=100, start=120, bottom=100, end=120)
    set_col_widths(table, [1.45, 5.2])

    doc.add_paragraph("")
    paragraph = doc.add_paragraph()
    paragraph.paragraph_format.space_before = Pt(50)
    run = paragraph.add_run("DECISION DOCUMENT")
    run.font.name = "Calibri"
    run.font.size = Pt(10)
    run.font.bold = True
    run.font.color.rgb = RGBColor.from_string(AMBER)
    paragraph = doc.add_paragraph()
    run = paragraph.add_run(
        "This report distinguishes the current implementation from the required production scope. "
        "Recommended catalog, pricing, policies and workflows are product decisions until formally approved."
    )
    run.font.name = "Calibri"
    run.font.size = Pt(10)
    run.font.color.rgb = RGBColor.from_string(MUTED)
    doc.add_page_break()


def add_document_control(doc):
    add_heading(doc, "Document Control", 1)
    add_table(
        doc,
        ["Field", "Value"],
        [
            ("Purpose", "Define the complete ecommerce product, required implementation, operating model and launch criteria."),
            ("Audience", "Founder, product owner, designer, developer, QA tester, content operator and admin team."),
            ("Current-state basis", "Observed PHP pages, API endpoints, database schema, current database records, design tokens and QA evidence."),
            ("Scope rule", "Current implementation and launch recommendations are explicitly separated to avoid treating assumptions as shipped behavior."),
            ("Version", "1.0 - 01 June 2026"),
        ],
        [1.55, 5.15],
    )
    add_heading(doc, "Status Legend", 2)
    add_status_key(doc)
    add_heading(doc, "How To Use This Report", 2)
    add_bullets(
        doc,
        [
            "Use Sections 2-5 to align the business, audience, product assortment and pricing.",
            "Use Sections 6-12 as the implementation scope for customer, admin, payment, email and content work.",
            "Use Sections 13-18 as the design, engineering, QA and launch acceptance reference.",
            "Resolve the open product decisions before final copywriting, catalog migration or production deployment.",
        ],
    )


def add_exec_summary(doc):
    add_heading(doc, "1. Executive Summary", 1)
    doc.add_paragraph(
        "UX Pacific Shop is an ecommerce storefront for UX learners, designers and early-career professionals. "
        "Its strongest product direction is a focused catalog of practical UX assets: templates, UI kits, mockups, "
        "research workbooks, interview resources, bundles, freebies and selected physical learning merchandise."
    )
    doc.add_paragraph(
        "The existing project already contains a meaningful end-to-end foundation: public discovery pages, authentication, "
        "cart, checkout, Razorpay integration, COD handling, customer account pages, wishlist, reviews, contact capture, "
        "admin operations and a MySQL schema. It is not production-ready yet. Payment lifecycle, order inventory handling, "
        "freebie downloads, canonical domain choice, placeholder content and mobile overflow require correction before launch."
    )
    add_callout(
        doc,
        "Recommended product position",
        "A curated UX resource shop that helps learners and working designers move faster: download useful assets immediately, "
        "buy guided physical practice tools when needed, and return to one account for orders, downloads, wishlists and reviews.",
    )
    add_heading(doc, "Launch Recommendation", 2)
    add_table(
        doc,
        ["Decision Area", "Recommendation", "Reason"],
        [
            ("Launch model", "Start with digital-first commerce and a small physical catalog.", "Digital products reduce fulfilment risk while the physical workflow is stabilized."),
            ("Core catalog", "Prioritize 8-12 polished digital SKUs, 4-5 bundles, 3-5 physical SKUs and 5-8 real freebies.", "Enough variety for discovery without diluting product quality."),
            ("Payments", "Razorpay Card/UPI plus COD for eligible physical products only.", "Matches the existing implementation and keeps the rule understandable."),
            ("Pricing", "Use simple INR price points and clearly separated personal/commercial licenses.", "Reduces checkout ambiguity and supports creator-friendly upsell."),
            ("Readiness", "Do not deploy to production until all P0 defects are resolved and retested.", "The current payment and inventory risks can create order or revenue errors."),
        ],
        [1.25, 2.55, 2.9],
    )


def add_vision_audience(doc):
    add_heading(doc, "2. Purpose, Goals and Audience", 1)
    add_heading(doc, "Why This Website Exists", 2)
    doc.add_paragraph(
        "UX Pacific Shop exists to convert useful UX learning and working materials into a dependable commerce experience. "
        "Instead of forcing customers to search across unrelated sources, the shop should present trusted, curated resources "
        "with clear previews, prices, license terms, delivery expectations and post-purchase access."
    )
    add_heading(doc, "Business Goals", 2)
    add_bullets(
        doc,
        [
            "Sell digital UX resources with immediate, reliable fulfilment after verified payment.",
            "Sell selected physical resources with accurate inventory, address capture, COD eligibility and shipping status.",
            "Build trust through a polished visual system, clear product information, transparent policies and dependable email communication.",
            "Create a reusable customer relationship through accounts, order history, wishlists, reviews, freebies and future offers.",
            "Give administrators one practical workspace for catalog, order, user, review, message and content operations.",
            "Measure conversion, cart health, payment success, product performance and repeat purchase behavior.",
        ],
    )
    add_heading(doc, "Primary Target Audiences", 2)
    add_table(
        doc,
        ["Audience", "Need", "Best Products", "Key UX Requirement"],
        [
            ("UX students and beginners", "Structured practice and portfolio support", "Workbooks, templates, starter kits, freebies", "Simple language, affordable entry prices and visible outcomes"),
            ("Early-career designers", "Job-readiness and presentation quality", "Portfolio kits, interview bundles, mockups, badge packs", "Fast comparison, clear previews and license clarity"),
            ("Freelancers and creators", "Reusable production assets", "UI kits, design systems, mockups, commercial licenses", "Commercial usage terms and efficient downloads"),
            ("Mentors and educators", "Teaching aids and repeatable exercises", "Printed workbooks, poster sets, bundles", "Bulk purchase path and clear fulfilment"),
            ("Returning customers", "Access and continuity", "Downloads, order details, wishlist, new releases", "Reliable account history and transactional email"),
        ],
        [1.3, 1.7, 1.8, 1.9],
        font_size=8.1,
    )
    add_heading(doc, "Core Personas", 2)
    add_table(
        doc,
        ["Persona", "Situation", "Purchase Trigger", "Likely Path"],
        [
            ("Aarav - UX learner", "Building a first case study", "Needs an affordable portfolio template and workbook", "Freebie -> product detail -> personal license -> Razorpay -> download"),
            ("Maya - freelance designer", "Preparing a client dashboard project", "Needs a UI kit with commercial usage", "Search -> UI kit -> commercial license -> checkout -> download"),
            ("Rohan - job seeker", "Preparing for interviews", "Needs bundled practice resources", "Bundles -> interview prep bundle -> cart -> payment -> order history"),
            ("Neha - design educator", "Runs a classroom workshop", "Needs printed workbooks and posters", "Physical catalog -> quantity -> address -> COD/online payment -> shipment updates"),
        ],
        [1.25, 1.9, 2.0, 1.55],
        font_size=8.1,
    )
    add_heading(doc, "Non-Goals For Initial Launch", 2)
    add_bullets(
        doc,
        [
            "Do not launch as a large general marketplace with third-party sellers.",
            "Do not add complex subscription tiers until one-time purchase, fulfilment and reporting are stable.",
            "Do not expand into international tax, multi-currency or multi-language checkout before the India-first baseline is reliable.",
            "Do not treat decorative content as a substitute for clear product previews, license terms and download instructions.",
        ],
    )


def add_catalog_pricing(doc):
    add_heading(doc, "3. Catalog Strategy, Categories and Pricing", 1)
    add_heading(doc, "Current Database Catalog", 2)
    doc.add_paragraph(
        "The database currently contains six categories, six active products, five active bundles and one active freebie. "
        "The table below records the observed state; it is not the recommended final launch catalog."
    )
    add_table(
        doc,
        ["Type", "Item", "Category", "Price", "Observed Notes"],
        [
            ("Product", "Design Career Badge Pack", "UX Resources", "INR 199", "Digital; old price INR 499"),
            ("Product", "Mobile App Mockup Pack", "Mockups", "INR 299", "Digital; featured"),
            ("Product", "SaaS UI Kit Essentials", "UI Kits", "INR 599", "Digital; featured"),
            ("Product", "UX Research Workbook", "Workbooks", "INR 499", "Digital; featured"),
            ("Product", "UXPacific UI Template", "Templates", "INR 399", "Digital; featured"),
            ("Product", "kmlm Booklet", "Templates", "INR 89", "Physical placeholder content; retire or replace"),
            ("Bundle", "Portfolio Builder Kit", "Bundles", "INR 999", "Observed active bundle"),
            ("Bundle", "UX Interview Prep Bundle", "Bundles", "INR 899", "Observed active bundle"),
            ("Bundle", "SaaS Launch Bundle", "Bundles", "INR 1,299", "Observed active bundle"),
            ("Bundle", "Research Sprint Bundle", "Bundles", "INR 799", "Observed active bundle"),
            ("Bundle", "Creator Starter Bundle", "Bundles", "INR 1,099", "Observed active bundle"),
            ("Freebie", "ui kits essentials", "UI Kit", "Free", "Placeholder file URL; replace with owned downloadable asset"),
        ],
        [0.7, 1.8, 1.1, 0.8, 2.3],
        font_size=7.9,
    )
    add_heading(doc, "Recommended Launch Categories", 2)
    add_table(
        doc,
        ["Category", "Purpose", "Suggested Range", "Launch Minimum"],
        [
            ("Templates", "Portfolio, case study, resume and presentation accelerators", "INR 299-799", "3 polished products"),
            ("UI Kits", "Reusable interface components and design-system starters", "INR 599-1,199", "2 polished products"),
            ("Mockups", "Presentation-ready screens, devices and portfolio visuals", "INR 299-699", "2 polished products"),
            ("UX Resources", "Interview, career, research and process assets", "INR 199-699", "3 polished products"),
            ("Workbooks", "Digital and printed guided practice material", "INR 349-699", "2 digital + 1 physical"),
            ("Bundles", "Outcome-focused combinations with visible savings", "INR 799-1,499", "4-5 bundles"),
            ("Freebies", "Lead-generation assets that demonstrate quality", "Free", "5-8 real downloads"),
        ],
        [1.05, 2.45, 1.2, 2.0],
    )
    add_heading(doc, "Recommended Digital Product Assortment", 2)
    add_table(
        doc,
        ["Product", "Category", "Personal", "Commercial", "Fulfilment"],
        [
            ("Portfolio Case Study Template", "Templates", "INR 399", "INR 599", "ZIP + PDF guide"),
            ("UX Resume and Cover Letter Kit", "Templates", "INR 299", "INR 449", "ZIP + editable files"),
            ("SaaS Dashboard UI Kit", "UI Kits", "INR 599", "INR 899", "Figma file + guide"),
            ("Design System Starter Pack", "UI Kits", "INR 799", "INR 1,199", "Figma file + tokens"),
            ("Mobile App Mockup Pack", "Mockups", "INR 299", "INR 499", "ZIP assets"),
            ("UX Research Workbook", "Workbooks", "INR 499", "INR 699", "PDF workbook"),
            ("UX Interview Toolkit", "UX Resources", "INR 399", "INR 599", "PDF + worksheets"),
            ("Design Career Badge Pack", "UX Resources", "INR 199", "INR 299", "ZIP assets"),
        ],
        [2.25, 1.1, 1.0, 1.0, 1.35],
    )
    add_heading(doc, "Recommended Physical Product Assortment", 2)
    add_table(
        doc,
        ["Product", "Category", "Price", "Fulfilment Note"],
        [
            ("UX Research Field Notebook", "Workbooks", "INR 349", "Inventory tracked; suitable entry product"),
            ("UX Canvas Workbook - Printed Edition", "Workbooks", "INR 599", "Address required; protect packaging"),
            ("UX Process Poster Set", "UX Resources", "INR 499", "Tube or flat-pack shipping decision required"),
            ("Creator Sticker Pack", "UX Resources", "INR 199", "Low-cost add-on and bundle enhancer"),
            ("Design Career Badge Pack - Physical", "UX Resources", "INR 249", "Optional companion to digital badge pack"),
        ],
        [2.8, 1.25, 0.9, 1.75],
    )
    add_heading(doc, "Bundle and Pricing Rules", 2)
    add_bullets(
        doc,
        [
            "Keep bundle savings visible and believable: target 20-35% below the sum of included products.",
            "Use rounded INR prices for customer-facing license variants. Avoid calculated-looking values such as INR 838.60.",
            "Show old price only when it reflects a genuine, supportable offer.",
            "Apply shipping to physical items only. Recommended baseline: INR 50 flat shipping, with an approved free-shipping threshold such as INR 999.",
            "Apply tax using one centrally managed server-side rule. The current checkout visibly applies 18%; finance approval is required before launch.",
            "Never accept product type, price, tax, stock or shipping values from the browser as authoritative.",
        ],
    )
    add_heading(doc, "Product Content Requirements", 2)
    add_table(
        doc,
        ["Field", "Required Standard"],
        [
            ("Title", "Clear outcome-focused name; remove placeholders and test data."),
            ("Slug", "Stable, lowercase, readable URL key."),
            ("Preview media", "Cover image plus useful gallery or preview pages; consistent aspect ratio."),
            ("Description", "What it solves, what is included, who it is for and how delivery works."),
            ("Type", "Digital, physical or both; always sourced from the database."),
            ("License", "Personal and commercial rules displayed before checkout."),
            ("Stock", "Real inventory for physical SKUs; digital products use availability status, not fake scarcity."),
            ("Files", "Owned, secure, versioned download assets; no placeholder URLs."),
            ("SEO", "Meta title, description, canonical URL and meaningful image alt text."),
        ],
        [1.2, 5.5],
    )


def add_information_architecture(doc, ia):
    add_heading(doc, "4. Information Architecture", 1)
    doc.add_paragraph(
        "The target information architecture keeps public discovery simple while separating authenticated customer tasks "
        "and administrative operations. The current PHP implementation already contains most of the listed destinations."
    )
    doc.add_picture(str(ia), width=Inches(6.7))
    add_heading(doc, "Current Sitemap and Purpose", 2)
    add_table(
        doc,
        ["Area", "Page / Route", "Purpose", "State"],
        [
            ("Discover", "index.php", "Homepage: hero, featured value, catalog entry points", "Implemented"),
            ("Discover", "shopAll.php", "All products, filters, sorting and pagination", "Implemented"),
            ("Discover", "category.php", "Category landing and filtered discovery", "Implemented"),
            ("Discover", "product.php?id={id}", "Product detail, license/type, quantity and add-to-cart", "Implemented"),
            ("Discover", "bundles.php", "Bundle catalog and bundle details", "Implemented"),
            ("Discover", "freebies.php", "Free downloadable lead magnets", "Partial"),
            ("Discover", "search.php?q={query}", "Search results", "Implemented"),
            ("Purchase", "cart.php", "Cart contents, quantities, removal and totals", "Implemented"),
            ("Purchase", "checkout.php", "Address, delivery, payment and order review", "Partial"),
            ("Purchase", "order-confirmation.php", "Post-order confirmation", "Partial"),
            ("Account", "signin.php / signup.php", "Authentication", "Implemented"),
            ("Account", "forgot-password.php / reset-password.php", "Password recovery", "Implemented"),
            ("Account", "account.php", "Profile, address and account overview", "Implemented"),
            ("Account", "orders.php", "Order history and review access", "Implemented"),
            ("Account", "wishlist.php", "Saved products", "Implemented"),
            ("Support", "contact.php", "Customer message submission", "Implemented"),
            ("Support", "policies.php", "Terms, privacy, delivery, return and download rules", "Implemented"),
            ("Admin", "admin/admin-login.php", "Administrator authentication", "Implemented"),
            ("Admin", "admin/admin-dashboard.php", "Back-office workspace", "Implemented"),
            ("Errors", "404.php / 500.php", "Custom error presentation", "Partial"),
        ],
        [0.75, 1.8, 3.25, 0.9],
        font_size=7.8,
    )
    add_heading(doc, "Recommended URL Direction", 2)
    add_table(
        doc,
        ["Current", "Recommended Public URL", "Reason"],
        [
            ("product.php?id=11", "/products/uxpacific-ui-template", "Readable, shareable and SEO-friendly product URL"),
            ("shopAll.php?category=ui-kits", "/categories/ui-kits", "Stable category landing"),
            ("bundles.php?quick=bundle&id=1", "/bundles/portfolio-builder-kit", "Clear bundle identity"),
            ("search.php?q=dashboard", "/search?q=dashboard", "Preserve search intent"),
            ("shopAll.php filters in JS only", "/shop?category=ui-kits&type=digital&sort=price-low", "Shareable filter state and browser history"),
        ],
        [2.15, 2.65, 2.0],
    )
    add_callout(
        doc,
        "Canonical-domain decision required",
        "Current project signals are inconsistent: the sitemap and robots file reference https://uxpacific.shop/ while application configuration has referenced https://shop.uxpacific.com. Choose one canonical production domain before launch and align redirects, sitemap, canonical tags, email links and payment callbacks.",
        color="FEF3C7",
    )


def add_click_redirects(doc):
    add_heading(doc, "5. Click-by-Click Navigation and Redirect Matrix", 1)
    doc.add_paragraph(
        "This matrix is the operational map for every major customer click. Where current behavior is incomplete, the required "
        "production behavior is recorded so development and QA have one acceptance reference."
    )
    add_heading(doc, "Public Discovery", 2)
    add_table(
        doc,
        ["Starting Location", "Click / Action", "Expected Destination or Result", "Notes"],
        [
            ("Header", "Logo", "index.php", "Returns to homepage"),
            ("Header", "Home", "index.php", "Primary navigation"),
            ("Header", "Products", "shopAll.php", "Product listing"),
            ("Header", "Bundles", "bundles.php", "Bundle listing"),
            ("Header", "Freebies", "freebies.php", "Lead-magnet listing"),
            ("Header", "Search icon / field submit", "search.php?q={encoded query}", "Empty-query behavior should remain helpful"),
            ("Header", "Cart icon", "cart.php", "Badge should reflect current quantity"),
            ("Header", "Sign In", "signin.php", "Hidden or replaced by account menu after login"),
            ("Header mobile", "Menu toggle", "Expandable mobile navigation", "Must not create horizontal overflow"),
            ("Homepage", "Featured product card", "product.php?id={product id}", "Use slug URL in future"),
            ("Homepage / categories", "Category card", "shopAll.php?category={slug}", "Prefer category route in future"),
            ("Shop listing", "Product card", "product.php?id={product id}", "Card metadata must match PDP"),
            ("Shop listing", "Category filter", "Filtered listing", "Required: reflect state in URL"),
            ("Shop listing", "Price range", "Filtered listing", "Required: reflect state in URL"),
            ("Shop listing", "Product-type checkbox", "Digital / physical subset", "Required: reflect state in URL"),
            ("Shop listing", "Sort selector", "Sorted visible products", "Required: reflect state in URL"),
            ("Product detail", "Choose license/type", "Price and eligibility refresh", "Server must validate type"),
            ("Product detail", "Quantity +/-", "Quantity updates within allowed limits", "Quantity cannot fall below 1"),
            ("Product detail", "Add to cart", "Cart state updates with visible feedback", "Cart should persist after refresh"),
            ("Bundles", "Open bundle", "Bundle detail modal or route", "Prefer stable shareable route"),
            ("Freebies", "Download", "File download or authentication gate", "Currently blocked by schema defect"),
        ],
        [1.25, 1.65, 2.4, 1.45],
        font_size=7.55,
    )
    add_heading(doc, "Authentication and Protected Areas", 2)
    add_table(
        doc,
        ["Starting Location", "Click / Action", "Expected Destination or Result", "Notes"],
        [
            ("signin.php", "Valid login", "Requested redirect or account.php", "Preserve checkout intent"),
            ("signin.php", "Invalid login", "Inline error; remain on sign-in", "Do not reveal whether an email exists"),
            ("signin.php", "Forgot password", "forgot-password.php", "Reset flow"),
            ("signup.php", "Valid registration", "Signed-in session and account or intended route", "Document exact product decision"),
            ("signup.php", "Already registered email", "Inline error with sign-in route", "No duplicate account"),
            ("forgot-password.php", "Submit email", "Neutral success response", "Send reset email if eligible"),
            ("reset-password.php", "Valid token + new password", "signin.php with success message", "Invalidate token after use"),
            ("Protected route", "Guest opens account.php", "signin.php?redirect=account.php", "Observed guard pattern"),
            ("Protected route", "Guest opens checkout.php", "signin.php?redirect=checkout.php", "Observed guard pattern"),
            ("Protected route", "Guest opens orders.php", "signin.php?redirect=orders.php", "Observed guard pattern"),
            ("Protected route", "Guest opens wishlist.php", "signin.php?redirect=wishlist.php", "Observed guard pattern"),
            ("Profile menu", "Logout", "Public state; redirect to homepage or sign-in", "Back button must not restore private content"),
            ("Google button", "Click Google sign in/up", "Real OAuth flow or remove button", "Currently placeholder toast only"),
        ],
        [1.25, 1.75, 2.5, 1.25],
        font_size=7.65,
    )
    add_heading(doc, "Cart, Checkout and Order", 2)
    add_table(
        doc,
        ["Starting Location", "Click / Action", "Expected Destination or Result", "Notes"],
        [
            ("cart.php", "Increase quantity", "Line subtotal and totals recalculate", "Server authoritative"),
            ("cart.php", "Decrease quantity", "Quantity stops at 1", "Remove is separate action"),
            ("cart.php", "Remove item", "Item removed; totals update", "Show empty state if final item"),
            ("cart.php", "Continue shopping", "shopAll.php", "Keep cart contents"),
            ("cart.php", "Proceed to checkout", "checkout.php or sign-in redirect", "Guest checkout not currently baseline"),
            ("checkout.php", "Digital-only cart", "Digital delivery notice; no address requirement", "COD unavailable"),
            ("checkout.php", "Physical cart", "Shipping address required", "INR 50 current shipping rule"),
            ("checkout.php", "Select Card / UPI", "Razorpay initiation", "Payment amount must equal server total"),
            ("checkout.php", "Select COD", "Allowed only for eligible physical order", "Digital order must reject COD"),
            ("checkout.php", "Place order", "Create valid order exactly once", "Prevent duplicate submission"),
            ("Payment gateway", "Payment success", "Verify payment -> order confirmation", "Commit stock after verification"),
            ("Payment gateway", "Payment failure/cancel", "Return actionable state; preserve cart", "Release any reservation"),
            ("order-confirmation.php", "View orders", "orders.php", "Do not trust localStorage as order source"),
            ("orders.php", "Open order detail", "Order detail UI", "Show fulfilment and payment status"),
            ("orders.php", "Submit eligible review", "Review recorded for moderation", "Prevent duplicates"),
        ],
        [1.2, 1.65, 2.45, 1.45],
        font_size=7.55,
    )
    add_heading(doc, "Admin Operations", 2)
    add_table(
        doc,
        ["Admin Location", "Action", "Expected Result"],
        [
            ("admin/admin-login.php", "Valid admin login", "Open admin dashboard"),
            ("Dashboard overview", "Open analytics", "Display operational metrics from server data"),
            ("Products", "Create / edit / duplicate / toggle / delete", "Catalog changes persist with validation and audit visibility"),
            ("Categories", "Create / edit / delete", "Category navigation and catalog remain consistent"),
            ("Bundles", "Create / edit / delete", "Bundle composition, price and visibility persist"),
            ("Freebies", "Create / edit / delete", "Owned file URL and download data persist"),
            ("Orders", "Filter / inspect / update status", "Customer-visible status stays synchronized"),
            ("Users", "Filter / inspect / block / edit", "Access change is explicit and traceable"),
            ("Reviews", "Moderate review", "Visibility changes after moderation"),
            ("Messages", "Read / update / archive", "Support inbox remains manageable"),
            ("Theme", "Toggle admin theme", "Preference persists without layout regression"),
        ],
        [1.65, 2.3, 3.0],
    )


def add_features(doc):
    add_heading(doc, "6. Complete Feature Inventory and Implementation Scope", 1)
    add_heading(doc, "Storefront Features", 2)
    add_table(
        doc,
        ["Feature", "Current State", "Launch Requirement"],
        [
            ("Homepage", "Implemented", "Clear positioning, featured products, categories, bundles, freebies and trust messaging"),
            ("Product listing", "Implemented", "Stable filters, sorting, pagination, empty state and URL-backed filter state"),
            ("Category discovery", "Implemented", "Readable category routes and consistent product counts"),
            ("Product detail", "Implemented", "Gallery, pricing, license/type, stock, quantity, add-to-cart, related content and SEO metadata"),
            ("Bundles", "Implemented", "Visible included products, savings, eligibility and stable URLs"),
            ("Freebies", "Partial", "Real owned files, download tracking, clear access rules and working schema"),
            ("Search", "Implemented", "Live suggestions, result page, no-result state, keyboard navigation and query persistence"),
            ("Policies", "Implemented", "Approved terms, privacy, refund, shipping and digital-delivery wording"),
            ("Contact", "Implemented", "Validated submission, visible success/failure and admin inbox workflow"),
            ("Error pages", "Partial", "Wire Apache routing to branded 404 and 500 pages"),
        ],
        [1.35, 1.0, 4.6],
    )
    add_heading(doc, "Authentication and Account Features", 2)
    add_table(
        doc,
        ["Feature", "Current State", "Launch Requirement"],
        [
            ("Sign up", "Implemented", "Validation, duplicate handling, password visibility, confirmation and mobile layout"),
            ("Sign in", "Implemented", "Valid/invalid behavior, remember-session decision, redirect intent and throttling"),
            ("Google OAuth", "Placeholder", "Implement completely or remove visible buttons"),
            ("Password reset", "Implemented", "SMTP delivery, expiring token, neutral response and retest"),
            ("Logout", "Implemented", "Session invalidation, redirect, back-button and multi-tab testing"),
            ("Profile", "Implemented", "Editable identity data with clear success/error feedback"),
            ("Password change", "Implemented", "Current-password verification and session policy"),
            ("Addresses", "Implemented", "Add, update, delete, default address and checkout reuse"),
            ("Order history", "Implemented", "Reliable server-backed detail, payment state, fulfilment state and download access"),
            ("Wishlist", "Implemented", "Add/remove/toggle, empty state and account visibility"),
            ("Account deletion", "Implemented API", "Confirmation, policy decision and safe data retention handling"),
        ],
        [1.45, 1.0, 4.5],
    )
    add_heading(doc, "Commerce Features", 2)
    add_table(
        doc,
        ["Feature", "Current State", "Launch Requirement"],
        [
            ("Cart", "Implemented", "Add, update, remove, merge, persistence and accurate totals"),
            ("Digital purchase", "Partial", "Verified payment, secure delivery entitlement and email"),
            ("Physical purchase", "Partial", "Real inventory, shipping address, stock lifecycle and fulfilment tracking"),
            ("Mixed cart", "Needs explicit rule", "Decide whether one order can contain both types and define shipping/COD behavior"),
            ("COD", "Implemented", "Physical-only server enforcement"),
            ("Razorpay Card/UPI", "Partial", "Create, verify, callback, webhook secret and failure recovery"),
            ("Coupons", "Required", "Admin-defined code, dates, eligibility, usage limits and server-calculated discount"),
            ("Tax", "Implemented visibly", "Finance-approved rule, server authority and invoice wording"),
            ("Shipping", "Implemented baseline", "Admin-configurable flat rate and approved free-shipping threshold"),
            ("Order confirmation", "Partial", "Server-backed confirmation page and reliable email"),
            ("Downloads", "Required hardening", "Authenticated entitlement, expiring links and versioned files"),
            ("Reviews", "Implemented", "Eligibility, duplicate prevention and moderation"),
        ],
        [1.45, 1.0, 4.5],
    )
    add_heading(doc, "Growth and Operations Features", 2)
    add_table(
        doc,
        ["Feature", "Priority", "Requirement"],
        [
            ("Analytics dashboard", "Launch", "Revenue, order, payment, product and fulfilment visibility"),
            ("Email templates", "Launch", "Brand-consistent reset, order, payment and delivery messages"),
            ("Coupon management", "Launch", "Needed for promotions without manual price changes"),
            ("SEO controls", "Launch", "Canonical tags, sitemap alignment, metadata and structured data"),
            ("Newsletter capture", "Next sprint", "Consent-aware subscriber capture and source tracking"),
            ("Abandoned cart reminders", "Future", "Only after consent, deliverability and event tracking exist"),
            ("Bulk orders", "Future", "Educator inquiry workflow for physical items"),
            ("Product versioning", "Future", "Notify digital buyers when updated files are available"),
            ("Download limits", "Future", "Configurable policy for fraud prevention without harming legitimate buyers"),
        ],
        [1.85, 1.0, 4.1],
    )


def add_flows(doc, purchase, lifecycle):
    add_heading(doc, "7. Digital and Physical Purchase Workflows", 1)
    doc.add_picture(str(purchase), width=Inches(6.7))
    add_heading(doc, "Digital Product Workflow", 2)
    add_bullets(
        doc,
        [
            "Customer discovers a digital product through homepage, listing, category, search or bundle.",
            "Product detail explains contents, preview media, license options, price and immediate delivery.",
            "Customer selects personal or commercial license and adds the product to cart.",
            "Checkout removes physical-only requirements and disables COD.",
            "Server recalculates product eligibility, price, tax and total using database values.",
            "Razorpay payment starts. On verified success, the order becomes paid exactly once.",
            "Customer reaches a server-backed confirmation page and receives an order email.",
            "Account order detail exposes the entitled download through a secure delivery mechanism.",
        ],
    )
    add_heading(doc, "Physical Product Workflow", 2)
    add_bullets(
        doc,
        [
            "Customer discovers a physical product with visible stock and delivery expectations.",
            "Product detail allows valid quantity selection and adds the physical SKU to cart.",
            "Checkout requires a valid shipping address and applies the approved shipping rule.",
            "Customer selects online payment or COD if eligible.",
            "Server validates inventory and creates the order. Inventory is committed only at the approved lifecycle event.",
            "Admin moves the order through processing, shipped and delivered states; tracking data should be added when available.",
            "Customer sees current fulfilment status in order history and receives email updates.",
        ],
    )
    add_heading(doc, "Mixed-Cart Product Decision", 2)
    add_callout(
        doc,
        "Required before implementation sign-off",
        "Decide whether customers may purchase digital and physical items together. Recommended baseline: allow a mixed cart, require address and shipping for the physical portion, disable COD if business policy requires online payment for mixed orders, and deliver digital entitlements only after the order payment condition is satisfied.",
        color="FEF3C7",
    )
    add_heading(doc, "Recommended Order Lifecycle", 2)
    doc.add_picture(str(lifecycle), width=Inches(6.7))
    add_heading(doc, "Lifecycle Rules", 2)
    add_bullets(
        doc,
        [
            "A draft or awaiting-payment order must never permanently consume stock.",
            "Online inventory movement is committed after verified payment; failed or cancelled payment releases any reservation.",
            "COD inventory policy must be explicit. Recommended: commit stock when COD order is confirmed and record cancellation reversals.",
            "Every status change should be visible to administrators and reflected consistently in the customer account.",
            "Digital download entitlements are issued only after the required payment state is reached.",
            "Webhook processing must be idempotent so repeated gateway events cannot duplicate an order or stock movement.",
        ],
    )


def add_architecture(doc, architecture):
    add_heading(doc, "8. Technical Architecture and Data Model", 1)
    doc.add_picture(str(architecture), width=Inches(6.7))
    add_heading(doc, "Current Technical Shape", 2)
    add_table(
        doc,
        ["Layer", "Current Implementation", "Production Requirement"],
        [
            ("Web", "Apache + PHP pages with CSS and JavaScript", "Harden routing, caching, canonical URLs and error handling"),
            ("API", "JSON endpoints grouped by auth, cart, catalog, order, payment, user and admin", "Consistent validation, CSRF posture, authorization and error contract"),
            ("Database", "MySQL schema plus migrations", "Apply required migrations, validate indexes, back up and rehearse restore"),
            ("Payments", "Razorpay create, verify and webhook endpoints", "Configure production keys and webhook secret; test success, fail, cancel and replay"),
            ("Email", "SMTP-based transactional mail", "Verify inbox delivery, templates, sender identity, logs and retry behavior"),
            ("Assets", "Local static assets and URLs", "Secure owned downloads, optimized images and stable file storage"),
        ],
        [1.0, 2.45, 3.45],
    )
    add_heading(doc, "Database Entity Catalog", 2)
    add_table(
        doc,
        ["Entity", "Purpose", "Key Relationships"],
        [
            ("categories", "Catalog grouping and navigation metadata", "products.category_id"),
            ("products", "Digital and physical SKU records", "category, cart, orders, bundle items, reviews, wishlist"),
            ("bundles", "Curated product combinations", "bundle_items"),
            ("bundle_items", "Bundle composition", "bundles, products"),
            ("freebies", "Free lead-magnet assets", "download tracking"),
            ("users", "Customer identity and profile", "cart, orders, addresses, wishlist, reviews, tokens"),
            ("addresses", "Reusable customer shipping addresses", "users"),
            ("cart", "Current intended purchase", "users/session, products"),
            ("orders", "Commerce transaction header", "users, order_items, payment and fulfilment data"),
            ("order_items", "Purchased product snapshots", "orders, products"),
            ("wishlist", "Saved customer products", "users, products"),
            ("reviews", "Product feedback and moderation", "users, products, orders"),
            ("contact_messages", "Support inquiries", "admin workflow"),
            ("password_reset_tokens", "Expiring recovery tokens", "users"),
            ("admins", "Back-office authentication", "admin workspace"),
            ("featured_items", "Homepage or merchandising promotion", "products or bundles"),
            ("inventory_movements", "Traceable stock changes", "products, orders"),
            ("user_tokens", "Session/token support where used", "users"),
        ],
        [1.55, 2.55, 2.8],
        font_size=8.0,
    )
    add_heading(doc, "Data Rules", 2)
    add_bullets(
        doc,
        [
            "Store an order-item snapshot of name, price, type and license at purchase time so later catalog edits do not rewrite history.",
            "Use database product records as the authority for price, availability, stock, shipping eligibility and type.",
            "Track stock changes with reason, order reference, quantity delta and timestamp.",
            "Store only the minimum customer information required for fulfilment and support.",
            "Keep payment provider IDs and verification state auditable without exposing secrets to the browser.",
        ],
    )


def add_api_catalog(doc):
    add_heading(doc, "9. API Catalog", 1)
    doc.add_paragraph(
        "The following API inventory is based on the current codebase. It should be treated as the regression-test map and authorization review checklist."
    )
    groups = [
        ("Authentication", [
            "_bootstrap.php", "auth/admin-login.php", "auth/check_admin.php", "auth/csrf.php",
            "auth/forgot-password.php", "auth/login.php", "auth/login-process.php", "auth/logout.php",
            "auth/reset-password.php", "auth/session.php", "auth/signup.php", "auth/verify-reset-token.php",
        ]),
        ("Cart, Catalog and Search", [
            "cart/add.php", "cart/list.php", "cart/merge.php", "cart/remove.php", "cart/update.php",
            "catalog/detail.php", "catalog/list.php", "catalog/view.php", "product/get_details.php", "product/search.php",
            "freebies/download.php",
        ]),
        ("Order, Payment and Customer", [
            "order/create.php", "order/get.php", "payment/razorpay-create-order.php",
            "payment/razorpay-verify.php", "payment/webhook.php", "contact/send.php",
            "reviews/submit.php", "user/delete_account.php", "user/profile.php",
            "user/update_password.php", "user/update_profile.php", "wishlist/add.php",
            "wishlist/list.php", "wishlist/remove.php", "wishlist/toggle.php",
        ]),
        ("Address", [
            "address/add.php", "address/delete.php", "address/get.php", "address/set-default.php", "address/update.php",
        ]),
        ("Admin", [
            "admin/stats/overview.php", "admin/stats/analytics.php", "admin/product/create.php",
            "admin/product/delete.php", "admin/product/duplicate.php", "admin/product/get.php",
            "admin/product/list.php", "admin/product/save.php", "admin/product/toggle_status.php",
            "admin/product/update.php", "admin/categories/list.php", "admin/categories/save.php",
            "admin/categories/delete.php", "admin/bundles/list.php", "admin/bundles/save.php",
            "admin/bundles/delete.php", "admin/freebies/list.php", "admin/freebies/save.php",
            "admin/freebies/delete.php", "admin/order/list.php", "admin/order/get_details.php",
            "admin/order/update_status.php", "admin/order/delete.php", "admin/user/list.php",
            "admin/user/save.php", "admin/user/block.php", "admin/user/delete.php",
            "admin/reviews/list.php", "admin/reviews/moderate.php", "admin/messages/list.php",
            "admin/messages/update.php",
        ]),
    ]
    rows = []
    for group, endpoints in groups:
        rows.append((group, "\n".join(endpoints)))
    add_table(doc, ["API Group", "Current Endpoints"], rows, [1.55, 5.25], font_size=7.75)
    add_heading(doc, "API Acceptance Rules", 2)
    add_bullets(
        doc,
        [
            "Require authentication and authorization for every private customer or admin endpoint.",
            "Validate CSRF protection strategy for state-changing browser requests.",
            "Return consistent JSON success, error and validation structures.",
            "Never reveal database errors, secrets, tokens or stack traces to the browser.",
            "Recalculate cart, checkout and payment amounts server-side.",
            "Make payment verification and webhook processing idempotent.",
            "Rate-limit authentication, password-reset, contact and freebie-download abuse paths.",
        ],
    )


def add_email(doc):
    add_heading(doc, "10. Email, Notification and Communication Plan", 1)
    doc.add_paragraph(
        "Email is part of the product, not a background detail. The current project contains SMTP configuration and the mail host was reachable during QA, "
        "but actual inbox delivery still needs end-to-end verification. A historical log entry recorded an SMTP expectation error, so production sign-off must include real mailbox evidence."
    )
    add_table(
        doc,
        ["Trigger", "Recipient", "Required Content", "Launch State"],
        [
            ("Password reset request", "Customer", "Secure expiring reset link, expiry note, ignore-if-not-requested guidance", "Required verification"),
            ("Order placed - online", "Customer", "Order ID, item summary, amount, payment state, support route", "Required verification"),
            ("Order placed - COD", "Customer", "Order ID, item summary, amount due, shipping address, support route", "Required verification"),
            ("Digital fulfilment", "Customer", "Order ID, download access instructions, license reminder", "Required"),
            ("Physical shipped", "Customer", "Order ID, carrier/tracking when available, expected delivery guidance", "Required"),
            ("Payment failed", "Customer", "Actionable retry route without duplicate order risk", "Required"),
            ("Contact submission", "Customer", "Acknowledgement and response expectation", "Recommended"),
            ("Contact submission", "Admin/support", "Message details and admin workflow link", "Recommended"),
            ("Low stock", "Admin", "Product, remaining inventory and action link", "Next sprint"),
        ],
        [1.55, 1.0, 3.55, 0.85],
        font_size=7.8,
    )
    add_heading(doc, "Email Verification Checklist", 2)
    add_bullets(
        doc,
        [
            "Send every launch-critical template to real test inboxes on at least two providers.",
            "Verify sender name, from address, reply-to address, subject, body, links and mobile rendering.",
            "Confirm reset links expire and cannot be reused.",
            "Confirm order emails are sent only after the correct order state transition.",
            "Check SMTP logs for errors, retry behavior and accidental duplicate sends.",
            "Verify no secret, internal path or sensitive payment payload is included.",
        ],
    )


def add_admin(doc):
    add_heading(doc, "11. Administration and Operating Model", 1)
    doc.add_paragraph(
        "The admin workspace is the operational control plane. The current dashboard already exposes the core sections below; launch readiness depends on permission checks, validation, reliable data and a repeatable operating process."
    )
    add_table(
        doc,
        ["Admin Module", "Current Capability", "Required Operating Standard"],
        [
            ("Overview", "Summary statistics and dashboard navigation", "Accurate server-derived KPIs and useful empty/error states"),
            ("Analytics", "Analytics view", "Revenue, order, payment, product and date-range reporting"),
            ("Products", "Create, edit, duplicate, toggle status, delete", "Validation, real assets, license/type rules and audit awareness"),
            ("Categories", "List, save, delete", "Slug uniqueness and safe deletion behavior"),
            ("Bundles", "List, save, delete", "Composition validation and savings calculation"),
            ("Freebies", "List, save, delete", "Working schema, owned files and download visibility"),
            ("Orders", "List, detail, status update, delete", "Status lifecycle, payment clarity and reversible-safe operations"),
            ("Users", "List, save, block, delete", "Least privilege, confirmation and retention policy"),
            ("Reviews", "List and moderate", "Clear visibility state and moderation workflow"),
            ("Messages", "List and update", "Support ownership, status and archive handling"),
            ("Theme", "Admin theme toggle", "Persistent, accessible light/dark appearance"),
        ],
        [1.2, 2.2, 3.8],
        font_size=7.9,
    )
    add_heading(doc, "Daily Admin Checklist", 2)
    add_bullets(
        doc,
        [
            "Review payment failures, awaiting-payment orders and COD orders.",
            "Progress physical orders through processing, shipped and delivered states.",
            "Monitor inventory and restock physical SKUs before they reach zero.",
            "Review customer messages and moderate new reviews.",
            "Check new product, bundle and freebie changes on the live storefront after publishing.",
            "Inspect mail and application logs for repeated errors.",
        ],
    )


def add_design_system(doc):
    add_heading(doc, "12. Design System", 1)
    doc.add_paragraph(
        "The current storefront already has a dark visual identity and recurring purple accent. The design system below turns that direction into a reusable product rulebook for customer and admin interfaces."
    )
    add_heading(doc, "Brand Direction", 2)
    add_table(
        doc,
        ["Attribute", "Rule"],
        [
            ("Position", "Confident, practical and creator-focused; premium enough to feel trustworthy without becoming intimidating."),
            ("Visual mood", "Dark digital workspace with purposeful purple accents, clear hierarchy and restrained glow."),
            ("Content tone", "Direct, encouraging and outcome-led. Explain what a customer gets and what it helps them do."),
            ("Imagery", "Show actual resource previews, not generic decoration. Keep product-cover ratios consistent."),
        ],
        [1.3, 5.4],
    )
    add_heading(doc, "Color Tokens", 2)
    add_table(
        doc,
        ["Token", "Observed / Recommended Value", "Use"],
        [
            ("bg-main", "#01010D", "Primary storefront background"),
            ("card-bg", "#050519", "Product cards and elevated surfaces"),
            ("card-bg-soft", "#07071D", "Secondary surface"),
            ("accent", "#6F4BFF", "Primary CTA, active state and focus emphasis"),
            ("accent-soft", "#9B8CFF", "Secondary accent and subtle highlight"),
            ("accent-hover", "#5B2EE6", "Interactive hover state"),
            ("text-primary", "#FFFFFF", "Primary dark-theme text"),
            ("text-muted", "rgba(255,255,255,0.75)", "Secondary dark-theme text"),
            ("success", "#10B981", "Confirmed state"),
            ("warning", "#F59E0B", "Attention state"),
            ("danger", "#EF4444", "Destructive and validation error state"),
            ("info", "#06B6D4", "Informational state"),
        ],
        [1.45, 2.05, 3.2],
    )
    add_heading(doc, "Typography", 2)
    add_table(
        doc,
        ["Role", "Recommendation", "Usage"],
        [
            ("Primary UI font", "Inter", "Navigation, forms, buttons, cards, tables and body copy"),
            ("Optional display font", "Gabarito or approved display face", "Marketing headings only; keep admin UI on Inter"),
            ("Body", "16px / 1.5", "Readable product and policy copy"),
            ("Small UI", "12-14px / 1.4", "Labels, badges and metadata"),
            ("Section heading", "28-40px", "Public page hierarchy"),
            ("Admin heading", "20-28px", "Operational hierarchy"),
        ],
        [1.55, 2.25, 2.9],
    )
    add_heading(doc, "Spacing, Radius and Elevation", 2)
    add_table(
        doc,
        ["System", "Tokens", "Rule"],
        [
            ("Spacing", "4, 8, 12, 16, 24, 32, 48, 64px", "Use consistent increments; avoid one-off spacing values unless layout demands it"),
            ("Radius", "6, 10, 14, 20px and pill", "Inputs/buttons 10px; cards 14-20px; badges pill"),
            ("Transitions", "150ms, 220ms, 340ms", "Use restrained motion for hover, modal and menu feedback"),
            ("Shadows", "sm, md, lg, xl", "Use elevation sparingly on overlays and cards"),
            ("Borders", "soft and subtle white alpha borders", "Separate dark surfaces without heavy outlines"),
        ],
        [1.1, 2.2, 3.4],
    )
    add_heading(doc, "Core Components", 2)
    add_table(
        doc,
        ["Component", "Required Variants and States"],
        [
            ("Button", "Primary, secondary, ghost, danger; default, hover, focus, disabled, loading"),
            ("Input", "Text, email, password, phone, textarea, select, checkbox, radio; default, focus, error, disabled"),
            ("Product card", "Digital, physical, bundle, freebie; price, old price, rating, stock/type badge, CTA"),
            ("Header", "Desktop nav, mobile menu, search, cart badge, signed-out and signed-in account states"),
            ("Modal", "Quick view, confirmation, review and support; focus trap, close control and mobile fit"),
            ("Toast", "Success, error, warning and info; accessible announcement and dismiss behavior"),
            ("Empty state", "Cart, wishlist, orders, search and admin tables with one clear next action"),
            ("Skeleton", "Listing, product detail, cart, orders and admin data loading"),
            ("Table", "Admin sort, filter, responsive overflow, empty and error state"),
            ("Badge", "Digital, physical, featured, bestseller, status and stock"),
        ],
        [1.35, 5.35],
    )
    add_heading(doc, "Responsive Rules", 2)
    add_table(
        doc,
        ["Viewport", "Required QA Sizes", "Design Rule"],
        [
            ("Mobile", "360, 375, 390, 414px", "Single-column priority, touch targets >=44px, no horizontal overflow, compact header"),
            ("Tablet", "768, 820, 1024px", "Adapt grid and filters deliberately; avoid stretched mobile layout"),
            ("Desktop", "1366, 1440, 1920px", "Constrain readable content width; preserve consistent card geometry"),
        ],
        [1.0, 2.0, 3.7],
    )
    add_heading(doc, "Accessibility Baseline", 2)
    add_bullets(
        doc,
        [
            "Provide meaningful alt text for product images and accessible labels for icon-only controls.",
            "Use visible keyboard focus states and logical tab order.",
            "Associate form errors with the relevant fields and announce important state changes.",
            "Meet readable contrast for text, disabled states and subtle borders.",
            "Use semantic headings and real buttons or links for interactive controls.",
            "Trap focus in modals and return focus to the triggering control on close.",
        ],
    )


def add_qa(doc):
    add_heading(doc, "13. Current QA Findings and Production Blockers", 1)
    doc.add_paragraph(
        "The findings below are known current-state risks discovered during QA inspection. They must be tracked independently and retested after correction. This report does not modify application code."
    )
    add_table(
        doc,
        ["ID", "Severity", "Area", "Finding", "Required Outcome"],
        [
            ("QA-001", "Critical", "Payment", "Razorpay webhook cannot operate without a configured webhook secret; endpoint returned service-unavailable behavior.", "Configure secret and verify signed webhook processing"),
            ("QA-002", "High", "Order security", "Order creation trusts browser-provided available_type, allowing shipping or inventory behavior to be influenced client-side.", "Derive product type from database only"),
            ("QA-003", "High", "Inventory", "Physical stock is decremented when an online-payment draft is created, before payment success; failure path does not restore it.", "Use reservation or post-verification commit"),
            ("QA-004", "High", "Freebies", "Freebie download fails because the active database schema lacks download_count expected by the endpoint.", "Apply schema migration and retest"),
            ("QA-005", "Medium", "Responsive CSS", "Homepage exhibits horizontal overflow at tested mobile widths.", "Remove overflow and retest 360-414px"),
            ("QA-006", "Medium", "Authentication UX", "Google sign-in/up buttons are visible but only show a coming-soon toast.", "Implement OAuth or remove CTA"),
            ("QA-007", "Medium", "Order confirmation", "Confirmation depends on localStorage, briefly shows success and redirects if lastOrder is absent; browser data is not authoritative.", "Load order by authorized server data"),
            ("QA-008", "Medium", "Content", "Placeholder physical product and freebie content are present in the live catalog data.", "Replace with approved real catalog content"),
            ("QA-009", "Medium", "Links", "Footer/community links include placeholder hash targets.", "Replace or remove dead CTAs"),
            ("QA-010", "Medium", "Errors", "Custom 404 page exists but unknown URLs use Apache default handling.", "Wire branded error routing"),
            ("QA-011", "Medium", "URL / UX", "Listing filters and sort state are not persisted in the URL.", "Use query-backed state"),
            ("QA-012", "Medium", "Security", "Environment secrets live under webroot, although blocked by .htaccess.", "Move secrets outside public document root where possible"),
            ("QA-013", "Low", "Footer", "Footer presentation and year/contact information are inconsistent.", "Normalize approved footer content"),
            ("QA-014", "Low", "Accessibility", "Shared navigation imagery includes missing or unclear alt behavior.", "Add meaningful accessible text"),
            ("QA-015", "Low", "Server hardening", "HTTP responses expose server-version detail.", "Review server banner policy"),
        ],
        [0.7, 0.7, 0.95, 3.35, 1.3],
        font_size=7.25,
    )
    add_heading(doc, "Known Working Baseline", 2)
    add_bullets(
        doc,
        [
            "Guest access guards redirect protected customer pages to sign-in with redirect intent.",
            "Basic invalid login, invalid signup and invalid forgot-password validation paths respond visibly.",
            "Cart quantity and visible tax totals recalculated correctly in sampled cart checks.",
            "Major sampled pages did not produce blocking JavaScript console errors.",
            "Debug scripts and sensitive dotfiles were blocked by web-server rules during inspection.",
            "SMTP host connectivity was reachable, but actual inbox delivery remains unverified.",
        ],
    )


def add_security_seo(doc):
    add_heading(doc, "14. Security, SEO and Performance Baseline", 1)
    add_heading(doc, "Security Checklist", 2)
    add_bullets(
        doc,
        [
            "Keep secrets outside public webroot where deployment allows; block dotfiles at the server level as defense in depth.",
            "Enforce server-side price, type, stock, shipping, tax, discount and order-total calculation.",
            "Protect admin routes and APIs with dedicated admin authorization checks.",
            "Use secure session cookie flags, rotate sessions after authentication and verify logout invalidation.",
            "Rate-limit login, password reset, contact and download endpoints.",
            "Sanitize output, validate inputs and prevent error detail leakage.",
            "Verify Razorpay signatures and make webhook/order handling idempotent.",
            "Use entitlement-based downloads instead of exposing unrestricted purchased-file URLs.",
            "Log important admin changes and stock movements.",
        ],
    )
    add_heading(doc, "SEO Checklist", 2)
    add_bullets(
        doc,
        [
            "Choose and enforce one canonical domain with HTTPS redirects.",
            "Add unique title, meta description and canonical tags to public content pages.",
            "Generate a sitemap from active public URLs and align robots.txt.",
            "Add product structured data where accurate, including price, availability and review aggregate only when valid.",
            "Use readable slug routes and redirect legacy query URLs if routes change.",
            "Avoid indexing account, cart, checkout, admin and internal API paths.",
        ],
    )
    add_heading(doc, "Performance Checklist", 2)
    add_bullets(
        doc,
        [
            "Optimize product media formats and dimensions; lazy-load below-the-fold images.",
            "Measure homepage, listing, product detail, cart and checkout on mobile network conditions.",
            "Reduce layout shift by reserving image and skeleton space.",
            "Cache safe static assets and compress CSS, JavaScript and images.",
            "Track API response time for search, cart, checkout, order creation and admin lists.",
            "Verify that third-party payment loading failure produces an actionable customer error.",
        ],
    )


def add_roadmap(doc):
    add_heading(doc, "15. Implementation Roadmap", 1)
    add_heading(doc, "Phase 0 - Production Blockers", 2)
    add_table(
        doc,
        ["Workstream", "Required Work", "Exit Criteria"],
        [
            ("Payment", "Configure webhook secret; verify signatures; test success, failure, cancel and replay", "No payment blocker; idempotent order behavior"),
            ("Order security", "Derive type, price, shipping and stock rules from database", "Browser tampering cannot alter server totals or fulfilment"),
            ("Inventory", "Replace draft-time permanent decrement with approved lifecycle", "Failure and cancellation preserve accurate stock"),
            ("Freebies", "Apply missing schema migration and replace placeholder file", "Download succeeds and count updates"),
            ("Content", "Remove placeholder catalog and dead links", "All public content approved"),
            ("Responsive", "Resolve homepage mobile overflow", "No horizontal scroll at required mobile widths"),
            ("Domain", "Approve canonical domain and align configuration", "All public links and callbacks use one domain"),
        ],
        [1.1, 3.95, 1.9],
    )
    add_heading(doc, "Phase 1 - Launch Baseline", 2)
    add_bullets(
        doc,
        [
            "Finalize launch catalog, license wording, shipping rule, tax decision and policy content.",
            "Harden digital download entitlement and order-confirmation data source.",
            "Verify password reset and all transactional emails with real inbox evidence.",
            "Add coupon management, robust empty/error/loading states and branded error routing.",
            "Complete responsive, browser, accessibility and regression passes across all major routes.",
            "Create backup, restore, deployment and rollback procedures.",
        ],
    )
    add_heading(doc, "Phase 2 - Conversion and Operations", 2)
    add_bullets(
        doc,
        [
            "Improve URL-backed filtering, SEO metadata and structured product data.",
            "Add operational analytics, inventory alerts and stronger admin auditability.",
            "Add newsletter capture with consent and campaign-source tracking.",
            "Improve product previews, related products and bundle upsell.",
        ],
    )
    add_heading(doc, "Phase 3 - Growth", 2)
    add_bullets(
        doc,
        [
            "Evaluate educator bulk-order inquiry, product update notifications and digital versioning.",
            "Evaluate abandoned-cart reminders only after consent and deliverability are established.",
            "Evaluate regional expansion, international payment and tax only after India-first stability.",
        ],
    )


def add_acceptance(doc):
    add_heading(doc, "16. Launch Acceptance Checklist", 1)
    rows = [
        ("Business", "Catalog, prices, licenses, tax, shipping and policies approved", "Required"),
        ("Content", "No placeholder copy, assets, URLs or dead links", "Required"),
        ("Authentication", "Signup, login, logout, reset, protected routes, persistence and mobile UI tested", "Required"),
        ("Catalog", "Listing, search, filters, sort, category, PDP, stock and out-of-stock tested", "Required"),
        ("Cart", "Add, duplicate, quantity, remove, persistence, tax, shipping and empty state tested", "Required"),
        ("Checkout", "Digital, physical, mixed-cart decision, address, COD and online payment tested", "Required"),
        ("Payment", "Success, fail, cancel, webhook and duplicate-event behavior verified", "Required"),
        ("Orders", "Confirmation, history, detail, status, stock movement and review eligibility tested", "Required"),
        ("Email", "Reset, order and fulfilment emails verified in real inboxes", "Required"),
        ("Admin", "Product, category, bundle, freebie, order, user, review and message operations tested", "Required"),
        ("Responsive", "360, 375, 390, 414, 768, 820, 1024, 1366, 1440 and 1920px checked", "Required"),
        ("Browser", "Chrome, Firefox, Edge and Safari-relevant checks completed", "Required"),
        ("Accessibility", "Keyboard, labels, focus, contrast, headings and modal behavior checked", "Required"),
        ("Security", "Authorization, server authority, secret handling, rate limits and error leakage reviewed", "Required"),
        ("SEO", "Canonical domain, redirects, sitemap, robots, metadata and index exclusions aligned", "Required"),
        ("Operations", "Backup, restore, deployment, rollback, monitoring and support ownership documented", "Required"),
    ]
    add_table(doc, ["Area", "Acceptance Check", "State"], rows, [1.15, 4.85, 1.0], font_size=8.0)


def add_kpis_decisions(doc):
    add_heading(doc, "17. KPIs and Open Decisions", 1)
    add_heading(doc, "Recommended KPIs", 2)
    add_table(
        doc,
        ["KPI", "Why It Matters"],
        [
            ("Store conversion rate", "Shows whether discovery and trust turn into purchases"),
            ("Product-detail to add-to-cart rate", "Measures product content and pricing effectiveness"),
            ("Cart-to-checkout completion rate", "Highlights checkout friction"),
            ("Payment success rate", "Detects gateway or verification failures"),
            ("Average order value", "Measures bundle and upsell strength"),
            ("Digital download success rate", "Confirms fulfilment reliability"),
            ("Physical fulfilment time", "Tracks operational quality"),
            ("Repeat purchase rate", "Measures long-term catalog value"),
            ("Freebie-to-purchase conversion", "Measures lead-magnet effectiveness"),
            ("Support contact rate per order", "Reveals unclear product, payment or fulfilment behavior"),
        ],
        [2.25, 4.75],
    )
    add_heading(doc, "Open Decisions Register", 2)
    add_table(
        doc,
        ["Decision", "Recommended Direction", "Owner"],
        [
            ("Canonical domain", "Choose uxpacific.shop or shop.uxpacific.com and standardize", "Founder / Engineering"),
            ("Mixed cart", "Allow with explicit shipping and COD rule", "Product / Operations"),
            ("Shipping threshold", "Approve INR 50 baseline and optional INR 999 free-shipping threshold", "Operations / Finance"),
            ("Tax rule", "Approve applicable tax and invoice wording", "Finance"),
            ("Digital license text", "Publish personal and commercial usage terms", "Product / Legal"),
            ("Refund policy", "Separate digital and physical policy wording", "Product / Legal"),
            ("Google OAuth", "Implement completely or remove public CTA", "Product / Engineering"),
            ("Physical launch SKUs", "Approve real items, packaging and stock quantities", "Operations"),
            ("Download delivery", "Choose secure entitlement and file-hosting approach", "Engineering"),
            ("Admin deletion policy", "Prefer archive/disable where history matters", "Product / Engineering"),
        ],
        [1.75, 4.1, 1.15],
        font_size=7.9,
    )


def add_appendix(doc):
    add_heading(doc, "18. Appendix: Route Reference and Test Matrix", 1)
    add_heading(doc, "Public and Customer Page Reference", 2)
    add_table(
        doc,
        ["Route", "Access", "Primary QA Coverage"],
        [
            ("index.php", "Public", "Hero, navigation, cards, footer, responsive overflow"),
            ("shopAll.php", "Public", "Catalog, category, type, price, sorting, pagination, URL state"),
            ("category.php", "Public", "Category loading and consistency"),
            ("product.php?id={id}", "Public", "Gallery, price, type/license, quantity, cart, stock"),
            ("bundles.php", "Public", "Bundle cards, details, contents and pricing"),
            ("freebies.php", "Public", "Download eligibility, file and tracking"),
            ("search.php?q={query}", "Public", "Valid, empty, invalid and no-result searches"),
            ("cart.php", "Public/session", "Add, update, remove, persistence, totals and empty state"),
            ("signin.php", "Public", "Valid, invalid, redirect, password visibility and mobile UI"),
            ("signup.php", "Public", "Required fields, duplicate email, passwords and mobile UI"),
            ("forgot-password.php", "Public", "Neutral response and email"),
            ("reset-password.php", "Public token", "Valid, invalid, expired and reused token"),
            ("checkout.php", "Protected", "Digital, physical, address, payment and duplicate prevention"),
            ("order-confirmation.php", "Order context", "Authorized server data and navigation"),
            ("account.php", "Protected", "Profile, password, addresses and links"),
            ("orders.php", "Protected", "Filter, details, status and reviews"),
            ("wishlist.php", "Protected", "Add, remove and empty state"),
            ("contact.php", "Public", "Validation, success and admin visibility"),
            ("policies.php", "Public", "Approved content and links"),
            ("404.php / 500.php", "Error", "Routing, branding and status code"),
        ],
        [1.8, 1.0, 4.2],
        font_size=7.8,
    )
    add_heading(doc, "Admin Page Reference", 2)
    add_table(
        doc,
        ["Route", "Access", "Primary QA Coverage"],
        [
            ("admin/admin-login.php", "Public admin entry", "Valid/invalid login and redirect"),
            ("admin/admin-dashboard.php", "Admin protected", "Every dashboard module, authorization and responsive table behavior"),
            ("admin/addproduct.php", "Admin protected", "Create validation, content and asset behavior"),
            ("admin/editproduct.php", "Admin protected", "Update, status and data persistence"),
        ],
        [2.15, 1.25, 3.6],
    )
    add_heading(doc, "QA Regression Viewports", 2)
    add_table(
        doc,
        ["Class", "Viewports", "Required Checks"],
        [
            ("Mobile", "360, 375, 390, 414px", "Header, menu, cards, PDP, cart, checkout, forms, modals, footer, overflow and touch targets"),
            ("Tablet", "768, 820, 1024px", "Grid, filters, account, checkout, admin tables and modal fit"),
            ("Desktop", "1366, 1440, 1920px", "Content width, grid density, header, cards, checkout and admin workspace"),
        ],
        [1.0, 1.9, 4.1],
    )
    add_heading(doc, "Final Project Definition", 2)
    doc.add_paragraph(
        "UX Pacific Shop should launch as a dependable, focused ecommerce product: customers discover high-quality UX resources, "
        "understand exactly what they are buying, complete an accurate payment flow, receive the right digital or physical fulfilment, "
        "and return to a useful account. Administrators should be able to operate the catalog and orders without manual database work. "
        "The next implementation effort should be governed by the blocker list and acceptance checklist in this report."
    )


def build():
    ia, purchase, architecture, lifecycle = build_diagrams()
    doc = setup_document()
    cover(doc)
    add_document_control(doc)
    add_exec_summary(doc)
    add_vision_audience(doc)
    add_catalog_pricing(doc)
    add_information_architecture(doc, ia)
    add_click_redirects(doc)
    add_features(doc)
    add_flows(doc, purchase, lifecycle)
    add_architecture(doc, architecture)
    add_api_catalog(doc)
    add_email(doc)
    add_admin(doc)
    add_design_system(doc)
    add_qa(doc)
    add_security_seo(doc)
    add_roadmap(doc)
    add_acceptance(doc)
    add_kpis_decisions(doc)
    add_appendix(doc)
    doc.core_properties.title = "UX Pacific Shop Master Project Report"
    doc.core_properties.subject = "Ecommerce product, feature, architecture, design system and launch blueprint"
    doc.core_properties.author = "Codex"
    doc.core_properties.keywords = "UX Pacific, ecommerce, product blueprint, design system, QA, architecture"
    doc.save(DOCX_PATH)
    print(DOCX_PATH)


if __name__ == "__main__":
    build()
